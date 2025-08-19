# File: ingest_pdf/pipeline/docx_ingest.py
# Path: {PROJECT_ROOT}\ingest_pdf\pipeline\docx_ingest.py
# Description:
#   Extracts and enriches concepts from DOCX files.

import json
import logging
from pathlib import Path
PROJECT_ROOT = Path(__file__).resolve().parents[1]
import sys

# Add parent directories to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from tools.enrich_concepts import ConceptEnricher

logger = logging.getLogger("docx_ingest")
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

VAULT_DIR = Path("data/memory_vault/memories")
VAULT_DIR.mkdir(parents=True, exist_ok=True)

enricher = ConceptEnricher()

def extract_concepts_from_docx(filepath: Path):
    """
    Extract concepts from DOCX files
    Uses python-docx to read content, then applies concept extraction
    """
    logger.info(f"📝 Extracting text from DOCX: {filepath}")
    
    try:
        from docx import Document
    except ImportError:
        logger.error("❌ python-docx not installed. Run: pip install python-docx")
        return []
    
    try:
        # Read DOCX content
        doc = Document(str(filepath))
        
        # Extract all text
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        combined_text = "\n".join(full_text)
        
        # Now extract concepts from the text
        # This is a simplified version - you might want to use your existing
        # concept extraction logic from ingest_pdf
        from ingest_pdf.extraction.concept_extraction import extract_semantic_concepts
        
        concepts = extract_semantic_concepts(combined_text)
        
        # Add source metadata
        for concept in concepts:
            if isinstance(concept, dict):
                if 'metadata' not in concept:
                    concept['metadata'] = {}
                concept['metadata']['source_file'] = str(filepath)
                concept['metadata']['source_type'] = 'docx'
        
        return concepts
        
    except Exception as e:
        logger.error(f"❌ Failed to extract from DOCX: {e}")
        # Return dummy concept for testing
        return [{
            "label": f"docx extract failed: {filepath.name}",
            "score": 0.1,
            "method": "docx_fallback",
            "metadata": {
                "error": str(e),
                "source_file": str(filepath)
            }
        }]

def ingest_docx_clean(filepath: Path):
    """Extract and enrich concepts from DOCX file"""
    logger.info(f"📄 Ingesting DOCX: {filepath.name}")
    
    # Extract concepts from DOCX
    concepts = extract_concepts_from_docx(filepath)
    
    enriched_count = 0
    for concept in concepts:
        # Enrich the concept
        enriched = enricher.enrich_concept(concept)
        
        # Write to vault
        _write_to_vault(enriched)
        enriched_count += 1
    
    logger.info(f"✅ DOCX ingestion complete: {enriched_count} concepts enriched and stored")
    return enriched_count

def _write_to_vault(concept):
    """Write enriched concept to memory vault"""
    label = concept.get("label", "unlabeled").replace(" ", "_")
    hash_id = abs(hash(label)) % (10 ** 8)
    timestamp = int(Path.cwd().stat().st_mtime)
    file_path = VAULT_DIR / f"tori_mem_{timestamp}_{hash_id:08x}.json"
    
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump({"content": concept}, f, indent=2)
    
    logger.debug(f"💾 Saved concept to: {file_path.name}")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Ingest DOCX file and extract concepts")
    parser.add_argument("filepath", type=Path, help="Path to DOCX file")
    
    args = parser.parse_args()
    
    if not args.filepath.exists():
        logger.error(f"❌ File not found: {args.filepath}")
        sys.exit(1)
    
    ingest_docx_clean(args.filepath)
